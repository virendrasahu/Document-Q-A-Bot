import os
from pathlib import Path
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from docx import Document

# Define data directory
try:
    from src.config import DATA_DIR
except ImportError:
    DATA_DIR = Path(__file__).resolve().parent / "data"

DATA_DIR.mkdir(parents=True, exist_ok=True)

def create_business_pdf():
    pdf_path = DATA_DIR / "business_doc.pdf"
    print(f"Creating business report PDF: {pdf_path}")
    doc = SimpleDocTemplate(str(pdf_path), pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
    
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontSize=24,
        leading=28,
        textColor='#1E3A8A', # Slate Blue
        alignment=TA_CENTER,
        spaceAfter=20
    )
    
    section_style = ParagraphStyle(
        'SectionHeader',
        parent=styles['Heading2'],
        fontSize=16,
        leading=20,
        textColor='#0D9488', # Teal
        spaceBefore=15,
        spaceAfter=10
    )
    
    body_style = ParagraphStyle(
        'BodyTextCustom',
        parent=styles['BodyText'],
        fontSize=10,
        leading=14,
        alignment=TA_JUSTIFY,
        spaceAfter=10
    )

    story = []

    # Page 1
    story.append(Paragraph("Acme Corp - Q1 2026 Earnings & Performance Report", title_style))
    story.append(Spacer(1, 20))
    
    story.append(Paragraph("1. Executive Summary", section_style))
    story.append(Paragraph(
        "Acme Corp had an exceptional first quarter of fiscal year 2026. The company successfully executed its "
        "strategic expansion initiatives, driving substantial growth across both software-as-a-service (SaaS) products "
        "and physical logistics offerings. Overall revenue reached record heights, supported by high customer retention "
        "and new enterprise agreements in Europe and the Asia-Pacific region.",
        body_style
    ))
    
    story.append(Paragraph("2. Financial Performance Metrics", section_style))
    story.append(Paragraph(
        "In Q1 2026, Acme Corp recorded a total revenue of $14.5 million. This represents a significant 15% increase "
        "year-over-year compared to Q1 2025, when revenue stood at $12.6 million. Net income for the quarter reached "
        "$2.1 million, showcasing a strong net margin of approximately 14.5%. Diluted earnings per share (EPS) was "
        "calculated at $0.42, exceeding consensus analyst estimates of $0.38.",
        body_style
    ))
    story.append(Paragraph(
        "The company's primary driver of growth was the newly launched CloudScale platform, which contributed "
        "$4.2 million in recurring subscription revenues during its first full quarter since public release.",
        body_style
    ))
    
    story.append(PageBreak())

    # Page 2
    story.append(Paragraph("3. Operational Highlights & Milestones", section_style))
    story.append(Paragraph(
        "A major operational milestone in Q1 2026 was the completion of our automated logistics hub in Frankfurt, "
        "Germany. This hub reduces delivery times across Central Europe by 35% and is expected to save the company "
        "approximately $450,000 annually in transport overheads. Additionally, our active enterprise customer base "
        "expanded by 84 new logos, bringing the total active enterprise subscriptions to 1,204 accounts.",
        body_style
    ))
    
    story.append(Paragraph("4. Forward-Looking Guidance", section_style))
    story.append(Paragraph(
        "Looking forward to Q2 2026, the management board projects revenues to fall in the range of $15.8 million "
        "to $16.2 million, representing a target midpoint growth of 17% year-over-year. The company plans to hire "
        "an additional 40 engineers to bolster the CloudScale product roadmap and prepare for the release of "
        "CloudScale Enterprise Plus in late Q3 2026. Capital expenditure is expected to remain stable at $1.2 million.",
        body_style
    ))
    
    doc.build(story)

def create_science_pdf():
    pdf_path = DATA_DIR / "science_paper.pdf"
    print(f"Creating science paper PDF: {pdf_path}")
    doc = SimpleDocTemplate(str(pdf_path), pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
    
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'PaperTitle',
        parent=styles['Heading1'],
        fontSize=20,
        leading=24,
        textColor='#111827', # Dark grey/black
        alignment=TA_CENTER,
        spaceAfter=15
    )
    
    author_style = ParagraphStyle(
        'AuthorText',
        parent=styles['Normal'],
        fontSize=11,
        leading=14,
        alignment=TA_CENTER,
        spaceAfter=20
    )
    
    section_style = ParagraphStyle(
        'SectionHeader',
        parent=styles['Heading2'],
        fontSize=14,
        leading=18,
        textColor='#374151',
        spaceBefore=12,
        spaceAfter=8
    )
    
    body_style = ParagraphStyle(
        'BodyTextCustom',
        parent=styles['BodyText'],
        fontSize=9.5,
        leading=13,
        alignment=TA_JUSTIFY,
        spaceAfter=10
    )

    story = []

    # Page 1
    story.append(Paragraph("Room-Temperature Superconductivity in Nitrogen-Doped Lutetium Hydride", title_style))
    story.append(Paragraph("Dr. Elena Vance, Department of Condensed Matter Physics, Horizon Institute of Science", author_style))
    story.append(Spacer(1, 10))
    
    story.append(Paragraph("Abstract", section_style))
    story.append(Paragraph(
        "We report evidence of room-temperature superconductivity in a nitrogen-doped lutetium hydride compound. "
        "By utilizing a nitrogen concentration of approximately 1.2 at.% and compressing the sample to a pressure "
        "of 10 kilobars (1.0 GPa), we observed a sharp drop in electrical resistance and a corresponding manifestation "
        "of the Meissner effect. The transition occurs at a critical temperature (Tc) of 294 Kelvin, which marks a "
        "historic milestone in solid-state physics and unlocks unprecedented possibilities for industrial applications.",
        body_style
    ))
    
    story.append(Paragraph("1. Introduction", section_style))
    story.append(Paragraph(
        "For over a century, superconductivity was restricted to extreme cryogenic environments. The discovery "
        "of copper-oxide high-temperature superconductors in the late 1980s pushed temperatures to 133 K, but required "
        "liquid nitrogen cooling. Recent advances in hydrogen-rich materials under massive pressures (above 150 GPa) "
        "opened pathways to near-room temperature transitions, but the crushing pressures rendered practical use "
        "impossible. Here, we present a synthesis pathway that achieves room-temperature superconductivity at a highly "
        "reduced pressure of 10 kbar.",
        body_style
    ))
    
    story.append(PageBreak())

    # Page 2
    story.append(Paragraph("2. Experimental Methodology", section_style))
    story.append(Paragraph(
        "A foil of high-purity lutetium (99.99%) was placed inside a diamond anvil cell along with a gaseous mixture of "
        "hydrogen (99%) and nitrogen (1%). The cell was heated to 330 degrees Celsius for 24 hours under a pressure "
        "of 2 GPa to facilitate compound synthesis. The resulting sample appeared as a vibrant blue compound. As "
        "the pressure was adjusted down to 10 kbar, the sample underwent a distinct color shift from blue to bright red. "
        "Resistance measurements were conducted using a standard four-probe configuration.",
        body_style
    ))
    
    story.append(Paragraph("3. Implications and Applications", section_style))
    story.append(Paragraph(
        "Achieving superconductivity at 294 Kelvin (approximately 21 degrees Celsius) and just 10 kbar of pressure "
        "removes the need for bulky liquid helium or nitrogen cooling loops. The red lutetium compound can be stabilized "
        "using thin-film deposition techniques on rigid substrates. Potential immediate applications include: "
        "(a) Lossless power grids that can transmit electricity across continents with zero ohmic losses; "
        "(b) Ultra-efficient compact fusion reactors utilizing high-temperature superconducting magnets; "
        "(c) Frictionless maglev trains operating on low-cost high-speed tracks.",
        body_style
    ))
    
    doc.build(story)

def create_factsheet_docx():
    docx_path = DATA_DIR / "factsheet.docx"
    print(f"Creating factsheet DOCX: {docx_path}")
    doc = Document()
    
    # Title
    doc.add_heading("BookExpert Platform Factsheet & FAQ", level=0)
    
    doc.add_heading("1. What is BookExpert?", level=1)
    doc.add_paragraph(
        "BookExpert is an advanced Document Q&A Bot built on a Retrieval-Augmented Generation (RAG) framework. "
        "It leverages Google's Gemini generative model and a localized ChromaDB vector store to search enterprise "
        "document libraries, extract precise answers, and cite sources inline. This prevents hallucinations by "
        "restricting the LLM to only answer queries using documents that have been ingested into the system."
    )
    
    doc.add_heading("2. Core Features", level=1)
    p = doc.add_paragraph()
    p.add_run("• Multi-format ingestion: ").bold = True
    p.add_run("Supports PDF, DOCX, and TXT files out of the box.\n")
    p.add_run("• Recursive chunking: ").bold = True
    p.add_run("Intelligently splits text at paragraph and sentence structures to maintain semantic clarity.\n")
    p.add_run("• Vector persistence: ").bold = True
    p.add_run("Embeds chunks via text-embedding-004 and saves them locally, avoiding repeated API token usage.\n")
    p.add_run("• Grounded answering: ").bold = True
    p.add_run("Uses a strict system prompt forcing the model to cite sources inline and return an error message if the information is unavailable.")
    
    doc.add_heading("3. Subscription Tiers & Pricing", level=1)
    doc.add_paragraph(
        "The BookExpert service is offered in three subscription tiers to cater to different user bases:"
    )
    
    p_free = doc.add_paragraph()
    p_free.add_run("• Free Tier ($0/month): ").bold = True
    p_free.add_run("Allows up to 100 queries per month. Connects to standard public endpoints. Supports documents up to 5 MB each.")
    
    p_pro = doc.add_paragraph()
    p_pro.add_run("• Professional Tier ($29/month): ").bold = True
    p_pro.add_run("Provides unlimited monthly queries, supports teams of up to 10 members, raises file limits to 50 MB, and offers priority processing.")
    
    p_ent = doc.add_paragraph()
    p_ent.add_run("• Enterprise Tier (Custom Pricing): ").bold = True
    p_ent.add_run("Deploys a dedicated instance on customer VPC, offers custom integrations, infinite file uploads, and SLA-backed support contracts.")
    
    doc.add_heading("4. Support & Technical Assistance", level=1)
    doc.add_paragraph(
        "For customer support, technical assistance, or custom enterprise inquiries, please contact our support team at "
        "support@bookexpert.ai or open a ticket in the admin console. Our technical documentation is available at "
        "https://docs.bookexpert.ai."
    )
    
    doc.save(str(docx_path))

def create_policy_pdf():
    pdf_path = DATA_DIR / "policy_doc.pdf"
    print(f"Creating policy PDF: {pdf_path}")
    doc = SimpleDocTemplate(str(pdf_path), pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
    
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'PolicyTitle',
        parent=styles['Heading1'],
        fontSize=20,
        leading=24,
        textColor='#7C2D12', # Warm orange/red-brown
        alignment=TA_CENTER,
        spaceAfter=15
    )
    
    section_style = ParagraphStyle(
        'SectionHeader',
        parent=styles['Heading2'],
        fontSize=14,
        leading=18,
        textColor='#451A03',
        spaceBefore=12,
        spaceAfter=8
    )
    
    body_style = ParagraphStyle(
        'BodyTextCustom',
        parent=styles['BodyText'],
        fontSize=10,
        leading=14,
        alignment=TA_JUSTIFY,
        spaceAfter=10
    )

    story = []

    # Page 1
    story.append(Paragraph("Acme Corp - Employee Remote Work & Equipment Policy", title_style))
    story.append(Spacer(1, 10))
    
    story.append(Paragraph("1. Eligibility Guidelines", section_style))
    story.append(Paragraph(
        "All full-time engineering, product, and design employees at Acme Corp are eligible for a hybrid or fully remote "
        "work arrangement, subject to manager approval. Employees must reside in a state or country where Acme Corp is "
        "registered as an employer. Fully remote employees are expected to maintain a stable broadband internet connection "
        "and a quiet working space suitable for video conferencing.",
        body_style
    ))
    
    story.append(Paragraph("2. Equipment Stipend & Allowances", section_style))
    story.append(Paragraph(
        "To support remote employees in setting up an ergonomic and productive home office, Acme Corp provides a "
        "one-time equipment stipend of up to $1,000. This stipend can be claimed via the Expensify portal within the "
        "first 90 days of employment. Approved items include ergonomic office chairs, desks (including standing desks), "
        "external computer monitors, keyboards, mice, and webcams. Computing hardware (laptops) is provided directly by "
        "the IT department and does not count towards the $1,000 stipend limit.",
        body_style
    ))
    
    story.append(PageBreak())

    # Page 2
    story.append(Paragraph("3. Core Collaboration Hours", section_style))
    story.append(Paragraph(
        "To facilitate communication across different time zones, Acme Corp has established core collaboration hours "
        "from 10:00 AM to 3:00 PM Eastern Standard Time (EST) / 7:00 AM to 12:00 PM Pacific Standard Time (PST). All remote "
        "employees are expected to be online and responsive on Slack and email during these core hours. Recurring team "
        "syncs, design reviews, and sprint planning sessions should be scheduled within this time block.",
        body_style
    ))
    
    story.append(Paragraph("4. Attendance & Offsites", section_style))
    story.append(Paragraph(
        "Remote employees are requested to attend the company's bi-annual in-person engineering offsite events. These "
        "events are typically held in June and January and last for four days. All travel expenses, lodging, and meals "
        "for official offsites are fully paid by the company in accordance with the corporate travel policy.",
        body_style
    ))
    
    doc.build(story)

def main():
    create_business_pdf()
    create_science_pdf()
    create_policy_pdf()
    create_factsheet_docx()
    print("All dummy documents generated successfully in the data folder.")

if __name__ == "__main__":
    main()
